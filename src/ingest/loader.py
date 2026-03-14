"""
loader.py — extrai texto de documentos para a base de conhecimento.
Lê de PDF_SOURCE_DIR (definido em .env). Nunca copia arquivos fonte.

Formatos suportados:
- PDF  → pymupdf4llm (Markdown) com fallback pdfplumber (texto plano)
- DOCX → python-docx → Markdown
- XLSX → openpyxl → Markdown (tabelas)
- HTML → markdownify → Markdown
- TXT/MD/CSV → leitura direta
"""

import csv
import io
import logging
import os
from dataclasses import dataclass
from pathlib import Path

import pdfplumber
from dotenv import load_dotenv

try:
    import pymupdf4llm
    _HAS_PYMUPDF4LLM = True
except ImportError:
    _HAS_PYMUPDF4LLM = False

try:
    import docx
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

try:
    import openpyxl
    _HAS_OPENPYXL = True
except ImportError:
    _HAS_OPENPYXL = False

try:
    import markdownify
    _HAS_MARKDOWNIFY = True
except ImportError:
    _HAS_MARKDOWNIFY = False


# Extensoes suportadas para upload
EXTENSOES_SUPORTADAS = {".pdf", ".docx", ".xlsx", ".html", ".htm", ".txt", ".md", ".csv"}

load_dotenv()

logger = logging.getLogger(__name__)

PDF_SOURCE_DIR = os.getenv("PDF_SOURCE_DIR", "")

# Mapeamento fixo: nome do arquivo → metadados da norma
NORMA_MAP: dict[str, dict] = {
    "EC132_2023.pdf": {
        "codigo": "EC132_2023",
        "nome": "Emenda Constitucional nº 132, de 20 de dezembro de 2023",
        "tipo": "EC",
        "numero": "132",
        "ano": 2023,
    },
    "LC214_2025.pdf": {
        "codigo": "LC214_2025",
        "nome": "Lei Complementar nº 214, de 16 de janeiro de 2025",
        "tipo": "LC",
        "numero": "214",
        "ano": 2025,
    },
    "LC227_2026.pdf": {
        "codigo": "LC227_2026",
        "nome": "Lei Complementar nº 227, de 2026",
        "tipo": "LC",
        "numero": "227",
        "ano": 2026,
    },
}


@dataclass
class DocumentoNorma:
    codigo: str
    nome: str
    tipo: str
    numero: str
    ano: int
    arquivo: str
    texto: str


def extrair_texto_pdf(caminho: Path) -> str:
    """
    Extrai texto de um PDF em formato Markdown via pymupdf4llm.

    Preserva tabelas, headers e estrutura hierárquica do documento.
    Fallback para pdfplumber (texto plano) se pymupdf4llm não estiver
    disponível ou falhar.
    """
    if _HAS_PYMUPDF4LLM:
        try:
            md_text = pymupdf4llm.to_markdown(str(caminho))
            if md_text and md_text.strip():
                logger.info("PDF extraído via pymupdf4llm (Markdown): %s", caminho.name)
                return md_text
            logger.warning("pymupdf4llm retornou vazio para %s, tentando pdfplumber", caminho.name)
        except Exception as e:
            logger.warning("pymupdf4llm falhou para %s: %s — fallback pdfplumber", caminho.name, e)

    # Fallback: pdfplumber (texto plano)
    paginas: list[str] = []
    with pdfplumber.open(caminho) as pdf:
        for i, pagina in enumerate(pdf.pages):
            texto = pagina.extract_text()
            if texto:
                paginas.append(texto)
            else:
                logger.debug("Página %d sem texto extraível em %s", i + 1, caminho.name)
    return "\n".join(paginas)


def extrair_texto_docx(caminho: Path) -> str:
    """Extrai texto de um DOCX e converte para Markdown."""
    if not _HAS_DOCX:
        raise ImportError("python-docx nao instalado. pip install python-docx")
    doc = docx.Document(str(caminho))
    partes: list[str] = []
    for para in doc.paragraphs:
        texto = para.text.strip()
        if not texto:
            continue
        # Mapear estilos de heading para Markdown
        if para.style and para.style.name:
            style = para.style.name.lower()
            if "heading 1" in style:
                partes.append(f"# {texto}")
            elif "heading 2" in style:
                partes.append(f"## {texto}")
            elif "heading 3" in style:
                partes.append(f"### {texto}")
            elif "heading 4" in style:
                partes.append(f"#### {texto}")
            else:
                partes.append(texto)
        else:
            partes.append(texto)

    # Extrair tabelas
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            # Header + separator + data
            header = rows[0]
            sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
            partes.append(header)
            partes.append(sep)
            partes.extend(rows[1:])
            partes.append("")

    logger.info("DOCX extraido: %s", caminho.name)
    return "\n\n".join(partes)


def extrair_texto_xlsx(caminho: Path) -> str:
    """Extrai planilha XLSX e converte para Markdown (tabelas)."""
    if not _HAS_OPENPYXL:
        raise ImportError("openpyxl nao instalado. pip install openpyxl")
    wb = openpyxl.load_workbook(str(caminho), read_only=True, data_only=True)
    partes: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        partes.append(f"## {sheet_name}")

        # Primeira linha como header
        header_cells = [str(c) if c is not None else "" for c in rows[0]]
        partes.append("| " + " | ".join(header_cells) + " |")
        partes.append("| " + " | ".join(["---"] * len(header_cells)) + " |")

        for row in rows[1:]:
            cells = [str(c) if c is not None else "" for c in row]
            # Ignorar linhas completamente vazias
            if all(c == "" or c == "None" for c in cells):
                continue
            partes.append("| " + " | ".join(cells) + " |")
        partes.append("")

    wb.close()
    logger.info("XLSX extraido (%d abas): %s", len(wb.sheetnames), caminho.name)
    return "\n".join(partes)


def extrair_texto_html(caminho: Path) -> str:
    """Converte HTML para Markdown."""
    if not _HAS_MARKDOWNIFY:
        raise ImportError("markdownify nao instalado. pip install markdownify")
    html_content = caminho.read_text(encoding="utf-8", errors="replace")
    md = markdownify.markdownify(html_content, heading_style="ATX", strip=["script", "style"])
    logger.info("HTML extraido: %s", caminho.name)
    return md


def extrair_texto_csv(caminho: Path) -> str:
    """Converte CSV para tabela Markdown."""
    conteudo = caminho.read_text(encoding="utf-8", errors="replace")
    reader = csv.reader(io.StringIO(conteudo))
    rows = list(reader)
    if not rows:
        return ""

    partes: list[str] = []
    header = rows[0]
    partes.append("| " + " | ".join(header) + " |")
    partes.append("| " + " | ".join(["---"] * len(header)) + " |")
    for row in rows[1:]:
        # Garantir mesmo numero de colunas
        cells = row + [""] * (len(header) - len(row))
        partes.append("| " + " | ".join(cells[:len(header)]) + " |")

    logger.info("CSV extraido (%d linhas): %s", len(rows), caminho.name)
    return "\n".join(partes)


def extrair_texto_plaintext(caminho: Path) -> str:
    """Le arquivo de texto plano (TXT, MD)."""
    texto = caminho.read_text(encoding="utf-8", errors="replace")
    logger.info("Texto plano extraido: %s", caminho.name)
    return texto


def extrair_texto(caminho: Path) -> str:
    """
    Extrai texto de qualquer formato suportado e retorna em Markdown.

    Formatos: PDF, DOCX, XLSX, HTML, TXT, MD, CSV.
    Raises ValueError para extensoes nao suportadas.
    """
    ext = caminho.suffix.lower()
    if ext == ".pdf":
        return extrair_texto_pdf(caminho)
    elif ext == ".docx":
        return extrair_texto_docx(caminho)
    elif ext == ".xlsx":
        return extrair_texto_xlsx(caminho)
    elif ext in (".html", ".htm"):
        return extrair_texto_html(caminho)
    elif ext == ".csv":
        return extrair_texto_csv(caminho)
    elif ext in (".txt", ".md"):
        return extrair_texto_plaintext(caminho)
    else:
        raise ValueError(
            f"Formato nao suportado: {ext}. "
            f"Formatos aceitos: {', '.join(sorted(EXTENSOES_SUPORTADAS))}"
        )


def extrair_texto_bytes(conteudo: bytes, filename: str) -> str:
    """
    Extrai texto a partir de bytes em memória (para uploads via API).

    Cria arquivo temporário, extrai, e limpa.
    """
    import tempfile
    ext = Path(filename).suffix.lower()
    if ext not in EXTENSOES_SUPORTADAS:
        raise ValueError(
            f"Formato nao suportado: {ext}. "
            f"Formatos aceitos: {', '.join(sorted(EXTENSOES_SUPORTADAS))}"
        )

    with tempfile.NamedTemporaryFile(suffix=ext, delete=True) as tmp:
        tmp.write(conteudo)
        tmp.flush()
        return extrair_texto(Path(tmp.name))


def carregar_normas() -> list[DocumentoNorma]:
    """
    Carrega todos os PDFs mapeados de PDF_SOURCE_DIR.
    Retorna lista de DocumentoNorma com texto completo.
    """
    if not PDF_SOURCE_DIR:
        raise EnvironmentError("PDF_SOURCE_DIR não definido no .env")

    source_dir = Path(PDF_SOURCE_DIR)
    if not source_dir.exists():
        raise FileNotFoundError(f"Diretório de PDFs não encontrado: {source_dir}")

    documentos: list[DocumentoNorma] = []
    for filename, meta in NORMA_MAP.items():
        caminho = source_dir / filename
        if not caminho.exists():
            logger.warning("PDF não encontrado, pulando: %s", caminho)
            continue

        logger.info("Carregando PDF: %s", filename)
        texto = extrair_texto_pdf(caminho)
        doc = DocumentoNorma(
            codigo=meta["codigo"],
            nome=meta["nome"],
            tipo=meta["tipo"],
            numero=meta["numero"],
            ano=meta["ano"],
            arquivo=str(caminho),
            texto=texto,
        )
        documentos.append(doc)
        logger.info("  → %d caracteres extraídos de %s", len(texto), filename)

    return documentos
